"""
Clean-text extraction from documents (PDF, DOCX, MD, TXT).

Strategy per file type:

    .md / .txt   → used as-is (nothing to convert)
    .docx        → paragraphs and tables via python-docx
    .pdf         → native text layer with PyMuPDF (instant and lossless); if the
                   PDF is scanned (no usable text) it falls back to local OCR with
                   DeepSeek-OCR.

OCR is imported lazily: loading torch/transformers takes several seconds, so the
cost is only paid when it is genuinely needed.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from core import console

logger = logging.getLogger(__name__)

TEXT_SUFFIXES = {".md", ".txt", ".mmd", ".markdown"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES

# Files that live in the data folders to explain them, not to be processed. They
# share the .md extension with real documents, so the extension alone cannot tell
# them apart: a README ends up offered as a leaflet, or sent to the model as a
# regulation, unless it is filtered out by name.
DOC_NAME_BLOCKLIST = {"readme", "leeme", "leame", "notes", "notas"}

# Below this average of characters per page, the PDF is assumed to be scanned.
MIN_CHARS_PER_PAGE = 120


def is_document(path: Path) -> bool:
    """Whether the file is a processable document and not folder documentation."""
    return (
        path.is_file()
        and path.suffix.lower() in SUPPORTED_SUFFIXES
        and not path.name.startswith(".")
        and path.stem.lower() not in DOC_NAME_BLOCKLIST
    )


def list_documents(directory: Path) -> list[Path]:
    """Processable documents in `directory`, sorted by name."""
    if not directory.is_dir():
        return []
    return sorted(p for p in directory.iterdir() if is_document(p))


@dataclass
class ExtractionResult:
    """Result of extracting text from a document."""

    text: str
    method: str
    source: Path
    pages: Optional[int] = None
    ocr_output_dir: Optional[Path] = None

    @property
    def char_count(self) -> int:
        return len(self.text)


def clean_text(raw: str) -> str:
    """
    Normalise the extracted text without altering its content.

    Unifies line breaks, drops form feeds and trailing spaces on each line, and
    collapses runs of 3 or more blank lines.
    """
    text = raw.replace("\r\n", "\n").replace("\r", "\n").replace("\f", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"[ \t]{3,}", "  ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


def extract_document(
    source: Path,
    ocr_output_dir: Optional[Path] = None,
    force_ocr: bool = False,
) -> ExtractionResult:
    """
    Extract a document's clean text, choosing the best strategy.

    Args:
        source: Input file (pdf / docx / md / txt).
        ocr_output_dir: OCR working folder (per-page images, .mmd). When omitted, a
            sibling folder of the PDF is used.
        force_ocr: Force OCR even when the PDF has a text layer.

    Returns:
        An ExtractionResult with the text and the method used.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the extension is not supported.
    """
    source = Path(source)
    if not source.exists():
        raise FileNotFoundError(f"Documento no encontrado: {source}")

    suffix = source.suffix.lower()

    if suffix in TEXT_SUFFIXES:
        console.info(f"{source.name}: ya es texto, no requiere conversión")
        return ExtractionResult(
            text=clean_text(source.read_text(encoding="utf-8-sig")),
            method="passthrough",
            source=source,
        )

    if suffix in DOCX_SUFFIXES:
        return _extract_docx(source)

    if suffix in PDF_SUFFIXES:
        return _extract_pdf(source, ocr_output_dir=ocr_output_dir, force_ocr=force_ocr)

    raise ValueError(
        f"Extensión no soportada: {source.suffix}. "
        f"Soportadas: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
    )


def extract_text(source: Path) -> str:
    """Shortcut returning just the text (used as step 1's `text_loader`)."""
    return extract_document(source).text


def _extract_docx(source: Path) -> ExtractionResult:
    """Extract paragraphs and tables from a DOCX."""
    from docx import Document

    document = Document(str(source))
    blocks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    console.info(f"{source.name}: {len(document.paragraphs)} párrafos extraídos del DOCX")
    return ExtractionResult(text=clean_text("\n".join(blocks)), method="docx", source=source)


def _extract_pdf(
    source: Path,
    ocr_output_dir: Optional[Path],
    force_ocr: bool,
) -> ExtractionResult:
    """Extract text from a PDF: native layer and, when that falls short, local OCR."""
    native_text = ""
    page_count = None

    if not force_ocr:
        try:
            import fitz  # PyMuPDF

            with fitz.open(str(source)) as document:
                page_count = document.page_count
                native_text = "\n\n".join(page.get_text("text") for page in document)
        except Exception as e:
            logger.warning(f"No se pudo leer la capa de texto de {source.name}: {e}")

        density = len(native_text.strip()) / max(page_count or 1, 1)
        if native_text.strip() and density >= MIN_CHARS_PER_PAGE:
            console.ok(
                f"{source.name}: texto nativo extraído "
                f"({page_count} páginas, {len(native_text)} caracteres)"
            )
            return ExtractionResult(
                text=clean_text(native_text),
                method="pdf_native",
                source=source,
                pages=page_count,
            )

        console.warn(
            f"{source.name}: sin capa de texto útil "
            f"({density:.0f} caracteres/página) → se usará OCR local"
        )
    else:
        console.info(f"{source.name}: OCR forzado por configuración")

    return _extract_pdf_with_ocr(source, ocr_output_dir=ocr_output_dir, page_count=page_count)


def _extract_pdf_with_ocr(
    source: Path,
    ocr_output_dir: Optional[Path],
    page_count: Optional[int],
) -> ExtractionResult:
    """Run the local OCR pipeline over the PDF and return the merged markdown."""
    # Lazy import: loading torch/transformers costs seconds.
    from etl.pdf_ocr_pipeline import PDFToOCRPipeline

    work_dir = Path(ocr_output_dir) if ocr_output_dir else source.parent / "ocr"
    console.info(f"Ejecutando OCR local sobre {source.name}")
    console.detail(f"salida de trabajo: {work_dir}")

    pipeline = PDFToOCRPipeline(pdf_path=source, output_base_dir=work_dir)
    result = pipeline.process()

    if not result.get("success"):
        raise RuntimeError(f"El OCR falló: {result.get('error')}")

    merged = Path(result["merged_markdown_file"])
    console.ok(f"OCR completado: {console.path_link(merged)}")

    return ExtractionResult(
        text=clean_text(merged.read_text(encoding="utf-8")),
        method="pdf_ocr",
        source=source,
        pages=result.get("summary", {}).get("total_pages", page_count),
        ocr_output_dir=Path(result["output_dir"]),
    )
