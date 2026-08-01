"""
One-shot adequation of the leaflet (step 4).

Takes step 3's compliance report, keeps only the rules in `missing` status from the
applicable dispositions, and generates the adequated version of the leaflet in ONE
single LLM call.

Outputs (in whatever directory the caller points at):
    prospecto_adecuado.json   full result (includes the filtered report)
    prospecto_adecuado.txt    clean, readable text
    prospecto_adecuado.docx   with the adequations highlighted in colour
"""
from __future__ import annotations

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Pt, RGBColor

from agents.llm_client import LONG_CALL_RETRIES, MAX_TIMEOUT, LLMClient
from agents.prompts import ADEQUATOR_INSTRUCTIONS
from core import console
from core.config import ModelConfig, settings

logger = logging.getLogger(__name__)

DEFAULT_BASE_NAME = "prospecto_adecuado"

# DOCX colours
_RED = RGBColor(255, 0, 0)      # adequations and "fill this in" markers
_GREEN = RGBColor(0, 128, 0)    # references {ref. {...}}
_BLACK = RGBColor(0, 0, 0)


def sanitize_json_string(json_str: str) -> str:
    """
    Drop the invalid control characters that break `json.loads`.

    Keeps tab, newline and carriage return; discards the rest of the 0-31 range.
    """
    control_chars = "".join(chr(i) for i in range(32) if i not in (9, 10, 13))
    return json_str.translate(str.maketrans("", "", control_chars))


def _strip_code_fences(text: str) -> str:
    """Remove the ```json ... ``` the model sometimes wraps the response in."""
    clean = text.strip()
    if clean.startswith("```json"):
        clean = clean[7:]
    elif clean.startswith("```"):
        clean = clean[3:]
    if clean.endswith("```"):
        clean = clean[:-3]
    return clean.strip()


class ProspectAdequator:
    """Adequates a leaflet against the unmet rules, in a single LLM query."""

    def __init__(self, config: Optional[ModelConfig] = None):
        """
        Args:
            config: Model parameters; defaults to `settings.adequator`.
        """
        self.config = config or settings.adequator
        # A single call that rewrites the whole leaflet: extra retries only
        # stretch the wait (see LONG_CALL_RETRIES).
        self.client = LLMClient.from_config(
            self.config, ADEQUATOR_INSTRUCTIONS, max_retries=LONG_CALL_RETRIES
        )
        logger.info(f"ProspectAdequator listo (model={self.config.model})")

    # ---- Input --------------------------------------------------------------

    @staticmethod
    def filter_compliance_report(compliance_report: Dict[str, Any]) -> Dict[str, Any]:
        """
        Reduce the report to what the adequator needs: applicable dispositions and,
        within them, only the rules whose status is `missing`.

        Returns:
            The filtered report, with a `summary` of totals.
        """
        applicable = [
            d for d in compliance_report.get("classification_results", [])
            if d.get("applies", False)
        ]
        details_by_id = {
            d.get("disposition_id"): d
            for d in compliance_report.get("compliance_details", [])
        }

        filtered_details = []
        for disposition in applicable:
            disposition_id = disposition.get("disposition_id")
            details = details_by_id.get(disposition_id)
            if not details:
                logger.warning(f"Sin detalles de compliance para {disposition_id}")
                continue

            missing_rules = [
                rule for rule in details.get("rules", [])
                if rule.get("status") == "missing"
            ]
            if not missing_rules:
                continue

            console.detail(f"{disposition_id}: {len(missing_rules)} reglas a adecuar")
            filtered_details.append({
                "disposition_id": disposition_id,
                "disposition_title": details.get("disposition_title"),
                "rules": missing_rules,
                "summary": {"total_rules": len(missing_rules), "missing": len(missing_rules)},
            })

        return {
            "prospect_filename": compliance_report.get("metadata", {}).get("prospect_name"),
            "timestamp": compliance_report.get("metadata", {}).get("timestamp"),
            "classification_results": applicable,
            "compliance_details": filtered_details,
            "summary": {
                "total_dispositions_applicable": len(applicable),
                "total_dispositions_with_missing_rules": len(filtered_details),
                "total_missing_rules": sum(d["summary"]["missing"] for d in filtered_details),
            },
        }

    # ---- LLM call -----------------------------------------------------------

    def adequate(
        self,
        prospect_text: str,
        filtered_compliance_report: Dict[str, Any],
        timeout: float = MAX_TIMEOUT,
    ) -> Dict[str, Any]:
        """
        Adequate the leaflet for ALL the missing rules in a single query.

        Args:
            prospect_text: Full text of the original leaflet.
            filtered_compliance_report: Filtered report (only `missing` rules).
            timeout: Maximum wait, in seconds (ceiling: MAX_TIMEOUT).

        Returns:
            The JSON returned by the model (includes `updated_prospect_text`).
        """
        summary = filtered_compliance_report["summary"]
        message_json = json.dumps(
            {"prospect_text": prospect_text, "compliance_report": filtered_compliance_report},
            indent=2,
            ensure_ascii=False,
        )

        console.info(
            f"Enviando al LLM: {summary['total_missing_rules']} reglas de "
            f"{summary['total_dispositions_with_missing_rules']} disposiciones"
        )
        console.detail(f"payload: {len(message_json)} caracteres · modelo: {self.config.model}")
        console.detail("esto puede tardar varios minutos…")
        logger.debug(f"Payload enviado al adecuador:\n{message_json}")

        response = self.client.run_text(message_json, json_mode=True, timeout=timeout)
        logger.debug(f"Respuesta del adecuador:\n{response}")

        try:
            return json.loads(sanitize_json_string(_strip_code_fences(response)))
        except json.JSONDecodeError:
            logger.error(f"Respuesta no parseable (primeros 1000 chars): {response[:1000]}")
            raise

    # ---- Outputs ------------------------------------------------------------

    @staticmethod
    def clean_text(updated_prospect_text: str) -> str:
        """
        Normalise the adequated text so it reads cleanly: unify line breaks, undo
        literal escape sequences and collapse surplus blank lines.
        """
        text = updated_prospect_text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\\n", "\n").replace("\\t", "\t")
        text = "\n".join(line.rstrip() for line in text.split("\n"))
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip() + "\n"

    def write_docx(self, text: str, docx_path: Path) -> Path:
        """
        Generate the DOCX of the adequated leaflet, interpreting the prompt's markers.

        Markers:
            **text**           → bold
            ╬ ... ╬            → adequation block delimiters (never printed)
            *{ref. {...}}*     → green italics (reference to a disposition/rule)
            *text*             → red italics (an adequation)
            [COMPLETAR ACÁ!]   → red italics
        """
        doc = Document()
        for section in doc.sections:
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

        lines = text.split("\n")
        inside_adequation_block = False

        for i, line in enumerate(lines):
            if line.strip() == "╬":
                inside_adequation_block = not inside_adequation_block
                continue

            if not line.strip():
                # Skip the blank line that precedes a reference.
                next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
                if next_line.startswith("*{ref.") and next_line.endswith("}*"):
                    continue
                doc.add_paragraph()
                continue

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            self._write_formatted_line(paragraph, line, inside_adequation_block)

        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(docx_path))
        return docx_path

    @staticmethod
    def _write_formatted_line(paragraph, line: str, inside_adequation_block: bool) -> None:
        """Split a line into runs and apply the matching formatting to each."""
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|\[COMPLETAR ACÁ!\])"

        for part in re.split(pattern, line):
            if not part:
                continue

            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                content = part[1:-1]
                run = paragraph.add_run(content)
                run.italic = True
                is_reference = content.strip().startswith("{ref.") and "}" in content
                run.font.color.rgb = _GREEN if is_reference else _RED
            elif part.startswith("[") and part.endswith("]"):
                run = paragraph.add_run(part)
                run.italic = True
                run.font.color.rgb = _RED
            else:
                run = paragraph.add_run(part)
                run.font.color.rgb = _RED if inside_adequation_block else _BLACK
            run.font.size = Pt(11)

    # ---- Orchestration ------------------------------------------------------

    def run(
        self,
        compliance_report_path: Path,
        prospect_path: Path,
        output_dir: Path,
        base_name: str = DEFAULT_BASE_NAME,
    ) -> Dict[str, Any]:
        """
        Run the whole adequation and write the three outputs.

        Args:
            compliance_report_path: JSON of the compliance report (step 3).
            prospect_path: Clean text of the leaflet (step 2).
            output_dir: Output folder (corridas/<timestamp>/documento-adecuado).
            base_name: Base name of the generated files.

        Returns:
            The result, with an `output_files` key pointing at what was generated.
        """
        compliance_report = json.loads(Path(compliance_report_path).read_text(encoding="utf-8-sig"))
        prospect_text = Path(prospect_path).read_text(encoding="utf-8-sig")
        console.ok(f"Reporte de compliance y prospecto cargados ({len(prospect_text)} caracteres)")

        filtered = self.filter_compliance_report(compliance_report)
        total_missing = filtered["summary"]["total_missing_rules"]

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        json_path = output_dir / f"{base_name}.json"
        txt_path = output_dir / f"{base_name}.txt"
        docx_path = output_dir / f"{base_name}.docx"

        result: Dict[str, Any] = {
            "prospect_filename": Path(prospect_path).name,
            "original_prospect_path": str(prospect_path),
            "compliance_report_path": str(compliance_report_path),
            "timestamp": datetime.now().isoformat(timespec="seconds"),
            "filtered_compliance_report": filtered,
            "summary": {
                "total_missing_rules": total_missing,
                "total_dispositions": filtered["summary"]["total_dispositions_with_missing_rules"],
                "adequation_performed": total_missing > 0,
            },
        }

        if total_missing == 0:
            console.ok("No hay reglas incumplidas: el prospecto ya es conforme")
            updated_text = prospect_text
            result["adequation_notes"] = "No se requirieron cambios: el prospecto ya es conforme"
        else:
            adequation_result = self.adequate(prospect_text, filtered)
            result["adequation_result"] = adequation_result
            updated_text = _extract_updated_text(adequation_result)

        output_files: Dict[str, str] = {}
        if updated_text:
            clean = self.clean_text(updated_text)
            txt_path.write_text(clean, encoding="utf-8")
            output_files["txt"] = str(txt_path)
            console.ok(f"Texto adecuado: {console.path_link(txt_path)}")
            try:
                self.write_docx(clean, docx_path)
                output_files["docx"] = str(docx_path)
                console.ok(f"DOCX adecuado: {console.path_link(docx_path)}")
            except Exception as e:
                logger.exception("No se pudo generar el DOCX")
                console.warn(f"No se pudo generar el DOCX: {e}")
        else:
            console.warn("El modelo no devolvió 'updated_prospect_text': no se generan TXT ni DOCX")

        output_files["json"] = str(json_path)
        result["output_files"] = output_files
        json_path.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
        console.ok(f"Resultado JSON: {console.path_link(json_path)}")

        return result


def _extract_updated_text(adequation_result: Any) -> Optional[str]:
    """Find `updated_prospect_text` in the response, tolerating one level of nesting."""
    if not isinstance(adequation_result, dict):
        return None
    updated = adequation_result.get("updated_prospect_text")
    if not updated and isinstance(adequation_result.get("result"), dict):
        updated = adequation_result["result"].get("updated_prospect_text")
    return updated
